"""Document conversion pipeline — the app's USP.

Any customer document in (Word / Excel / PDF / CSV / text, in their own words)
→ an import-ready, human-reviewed set of checkpoint answers out, plus a
converted Excel file the operator can check, edit, archive, or re-import.

Design: conversion runs in the BACKGROUND (a job with visible progress), so a
big document converting slowly is fine — nothing blocks the browser.

Order of attack per document:
  1. Deterministic pass — if the document has structure (a control column and a
     status column, like our template or most gap registers), it parses
     instantly and completely. No AI involved.
  2. Chunked assisted pass — free-form prose is split into passages; for each
     passage we shortlist the ~8 checkpoints whose keywords appear and ask the
     self-hosted model only about those. Slow but steady, progress per chunk,
     nothing leaves the environment.
  3. Human review — nothing is applied to the questionnaire until an operator
     approves each suggestion. The converted .xlsx is generated either way.
"""
from __future__ import annotations

import io
import json
import threading
import uuid
from datetime import datetime, timezone

from sqlalchemy import delete, select

from ..db import SessionLocal
from ..models import AiSuggestion, ImportJob

# How many deterministic rows we consider "the document was structured — done".
STRUCTURED_THRESHOLD = 3


# ---------------- text extraction (any supported format) ----------------

def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for row in t.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            return "\n".join((pg.extract_text() or "") for pg in PdfReader(io.BytesIO(data)).pages)
        if name.endswith((".xlsx", ".xlsm")):
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets:
                for r in ws.iter_rows(values_only=True):
                    rows.append(" | ".join("" if c is None else str(c) for c in r))
            return "\n".join(rows)
        return data.decode("utf-8", errors="replace")
    except Exception:
        return data.decode("utf-8", errors="replace")


# ---------------- the background job ----------------

def start_job(db, company_id: str, filename: str, data: bytes, actor_email: str) -> ImportJob:
    job = ImportJob(company_id=company_id, filename=filename[:250],
                    status="queued", stage="Queued", created_by=actor_email)
    db.add(job)
    db.commit()
    db.refresh(job)
    t = threading.Thread(target=_run_job, args=(job.id, filename, data), daemon=True)
    t.start()
    return job


def _set(job_id: str, **fields) -> None:
    """Progress updates use their own short-lived session so every step is
    visible to any web worker immediately."""
    with SessionLocal() as s:
        j = s.get(ImportJob, job_id)
        if j:
            for k, v in fields.items():
                setattr(j, k, v)
            s.commit()


def _store_suggestions(company_id: str, filename: str, items: list[dict]) -> None:
    with SessionLocal() as s:
        s.execute(delete(AiSuggestion).where(AiSuggestion.company_id == company_id))
        batch = str(uuid.uuid4())
        for it in items:
            s.add(AiSuggestion(company_id=company_id, batch=batch,
                               control_id=it["controlId"], status=it["status"],
                               evidence=(it.get("evidence") or "")[:1000],
                               source_quote=(it.get("sourceQuote") or "")[:400],
                               confidence=it.get("confidence") or "high",
                               source_name=filename[:250]))
        s.commit()


def _run_job(job_id: str, filename: str, data: bytes) -> None:
    try:
        _set(job_id, status="running", stage="Reading the document")
        text = extract_text(filename, data)
        if not (text or "").strip():
            _set(job_id, status="error", stage="Could not read the file",
                 note="No readable text found. If this is a scanned image, please share a "
                      "version with selectable text.",
                 finished_at=datetime.now(timezone.utc))
            return

        with SessionLocal() as s:
            from .rulebook_service import latest_rulebook
            rb = latest_rulebook(s)
        controls = rb["controls"]
        cats = {c["id"]: c["name"] for c in rb["categories"]}

        # ---- pass 1: deterministic (instant when the document is structured)
        _set(job_id, stage="Looking for structured rows (control + status columns)")
        from .import_parser import build_id_lookup, parse_upload
        lookup = build_id_lookup(c["id"] for c in controls)
        rows = parse_upload(filename, data, lookup)
        if len(rows) >= STRUCTURED_THRESHOLD:
            items = [{"controlId": r["controlId"], "status": r["status"],
                      "evidence": r.get("evidence", ""),
                      "sourceQuote": "", "confidence": "high"} for r in rows]
            _store_suggestions_job_done(job_id, filename, items,
                                        note=f"Structured document — parsed directly "
                                             f"({len(items)} answers, no AI needed).")
            return

        # ---- pass 2: chunked assisted conversion (free-form prose)
        from . import ai_mapper
        ok, why = ai_mapper.provider_available()
        if not ok:
            if rows:  # a few deterministic rows are still worth delivering
                items = [{"controlId": r["controlId"], "status": r["status"],
                          "evidence": r.get("evidence", ""), "sourceQuote": "",
                          "confidence": "high"} for r in rows]
                _store_suggestions_job_done(job_id, filename, items,
                                            note=f"Parsed {len(items)} structured row(s). "
                                                 f"The assisted reader is unavailable ({why})")
            else:
                _set(job_id, status="error", stage="Assisted reader unavailable", note=why,
                     finished_at=datetime.now(timezone.utc))
            return

        from .import_parser import match_control_id, normalize_status
        by_id = {c["id"]: c for c in controls}
        kw = ai_mapper._control_keywords(controls, cats)
        # Small passages, document order, NO skipping: a fast machine finishes
        # sooner, a slow machine takes longer — either way the whole document
        # gets converted. Small slices also read far better on a small model.
        passages = ai_mapper._chunks(text, size=700)
        _set(job_id, total_chunks=len(passages), stage="Converting passage by passage")

        best: dict = {r["controlId"]: {"controlId": r["controlId"], "status": r["status"],
                                       "evidence": r.get("evidence", ""), "sourceQuote": "",
                                       "confidence": "high"} for r in rows}
        for i, passage in enumerate(passages, start=1):
            candidates = ai_mapper._shortlist(passage, kw, by_id, top=8)
            if candidates:
                items = []
                for attempt in (1, 2):  # one retry — a hiccup must not drop a passage
                    try:
                        raw = ai_mapper._call_ollama(
                            ai_mapper._prompt(candidates, cats, passage), timeout=90)
                        items = ai_mapper._extract_items(json.loads(raw))
                        break
                    except Exception:
                        if attempt == 1:
                            import time
                            time.sleep(3)
                allowed = {c["id"] for c in candidates}
                for m in items if isinstance(items, list) else []:
                    if not isinstance(m, dict):
                        continue
                    cid = match_control_id(str(m.get("controlId") or m.get("id") or ""), lookup)
                    st = normalize_status(str(m.get("status") or ""))
                    if not cid or not st or cid not in allowed:
                        continue
                    cand = {"controlId": cid, "status": st,
                            "evidence": str(m.get("evidence", ""))[:400],
                            "sourceQuote": str(m.get("sourceQuote", "") or m.get("quote", ""))[:300],
                            "confidence": str(m.get("confidence", "")).lower()[:10] or "medium"}
                    prev = best.get(cid)
                    rank = {"high": 3, "medium": 2, "low": 1}
                    if not prev or rank.get(cand["confidence"], 0) > rank.get(prev["confidence"], 0):
                        best[cid] = cand
            _set(job_id, done_chunks=i, found=len(best),
                 stage=f"Converting passage {i} of {len(passages)}")

        items = sorted(best.values(), key=lambda x: x["controlId"])
        if items:
            _store_suggestions_job_done(job_id, filename, items,
                                        note=f"Converted {len(items)} checkpoint answer(s) from "
                                             f"{len(passages)} passage(s). Review before applying.")
        else:
            _set(job_id, status="done", found=0, finished_at=datetime.now(timezone.utc),
                 stage="Finished — nothing matched",
                 note="The document didn't contain wording that lines up with the DPDPA "
                      "checkpoints. The template or paste import is the reliable path for it.")
    except Exception as ex:  # never leave a job stuck in 'running'
        _set(job_id, status="error", stage="Conversion failed",
             note=f"{type(ex).__name__}: {ex}", finished_at=datetime.now(timezone.utc))


def _store_suggestions_job_done(job_id: str, filename: str, items: list[dict], note: str) -> None:
    with SessionLocal() as s:
        j = s.get(ImportJob, job_id)
        company_id = j.company_id
    _store_suggestions(company_id, filename, items)
    _set(job_id, status="done", found=len(items), note=note,
         stage="Finished", finished_at=datetime.now(timezone.utc))


# ---------------- the converted, re-usable file ----------------

def build_converted_workbook(rulebook: dict, brand: str, company_name: str,
                             filename: str, suggestions: list) -> bytes:
    """The deliverable: a clean Excel file of everything the conversion found.
    Column headers match the direct importer, so the file round-trips."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    cats = {c["id"]: c["name"] for c in rulebook.get("categories", [])}
    titles = {c["id"]: c["title"] for c in rulebook.get("controls", [])}
    sections = {c["id"]: c["category"] for c in rulebook.get("controls", [])}

    wb = Workbook()
    ws = wb.active
    ws.title = "Converted answers"
    ws["A1"] = f"{brand} — converted answers"
    ws["A1"].font = Font(size=15, bold=True, color="147A3D")
    ws["A2"] = (f"Company: {company_name}   ·   Source document: {filename}   ·   "
                f"Converted: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    ws["A3"] = ("Review each row, correct anything, then upload this file back via "
                "Direct import — or apply directly from the review screen in the app.")
    ws["A3"].font = Font(italic=True, size=10, color="55677A")

    headers = ["Control ID", "Section", "Checkpoint", "Status", "Evidence", "Confidence", "Source quote"]
    hrow = 5
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=hrow, column=i, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="13202E")
    widths = [12, 22, 46, 12, 50, 12, 40]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    for r, sug in enumerate(suggestions, start=hrow + 1):
        ws.cell(row=r, column=1, value=sug.control_id)
        ws.cell(row=r, column=2, value=cats.get(sections.get(sug.control_id, ""), ""))
        ws.cell(row=r, column=3, value=titles.get(sug.control_id, ""))
        ws.cell(row=r, column=4, value=sug.status)
        ws.cell(row=r, column=5, value=sug.evidence)
        ws.cell(row=r, column=6, value=sug.confidence)
        ws.cell(row=r, column=7, value=sug.source_quote)
        for col in (3, 5, 7):
            ws.cell(row=r, column=col).alignment = Alignment(wrap_text=True, vertical="top")

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()
