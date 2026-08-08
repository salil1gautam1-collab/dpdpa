"""Parse client compliance data from simple pasted text or an uploaded file
(CSV / Excel / Word / PDF / JSON) into normalized questionnaire answers.

The goal is to "make sense of simple data": an operator can paste lines like
    SEC-01, gap, security not compliant, IT
or upload the client's own data sheet, and we extract {controlId, status,
evidence, department} rows — fuzzy-matching control ids and mapping loose status
words to the five canonical statuses.
"""
from __future__ import annotations

import csv
import io
import json
import re

VALID_STATUS = {"COMPLIANT", "PARTIAL", "GAP", "NA", "TBC"}

# loose word -> canonical status
_STATUS_WORDS = {
    "compliant": "COMPLIANT", "compliance": "COMPLIANT", "yes": "COMPLIANT", "y": "COMPLIANT",
    "ok": "COMPLIANT", "pass": "COMPLIANT", "c": "COMPLIANT", "met": "COMPLIANT",
    "partial": "PARTIAL", "partially": "PARTIAL", "partially compliant": "PARTIAL", "p": "PARTIAL",
    "in progress": "PARTIAL", "wip": "PARTIAL",
    "gap": "GAP", "non compliant": "GAP", "non-compliant": "GAP", "not compliant": "GAP",
    "noncompliant": "GAP", "nc": "GAP", "no": "GAP", "n": "GAP", "fail": "GAP", "open": "GAP",
    "na": "NA", "n/a": "NA", "not applicable": "NA", "notapplicable": "NA",
    "tbc": "TBC", "to be confirmed": "TBC", "tobeconfirmed": "TBC", "pending": "TBC",
    "unknown": "TBC", "unsure": "TBC", "review": "TBC",
}

_HEADER_HINTS = {
    "id": {"controlid", "control", "id", "checkpoint", "ref", "control id", "control ref"},
    "status": {"status", "state", "result", "compliance", "answer", "position"},
    "evidence": {"evidence", "remark", "remarks", "note", "notes", "comment", "comments",
                 "finding", "details", "detail", "current position"},
    "department": {"department", "dept", "owner", "team", "function"},
}


def normalize_status(value: str) -> str | None:
    if not value:
        return None
    v = value.strip()
    up = v.upper()
    if up in VALID_STATUS:
        return up
    key = re.sub(r"\s+", " ", v.lower()).strip()
    if key in _STATUS_WORDS:
        return _STATUS_WORDS[key]
    key2 = re.sub(r"[^a-z/]", "", v.lower())
    return _STATUS_WORDS.get(key2)


def _norm_id(s: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]", "", s or "").upper()
    m = re.match(r"^([A-Z]+)0*(\d+)$", s)
    return f"{m.group(1)}{int(m.group(2))}" if m else s


def build_id_lookup(valid_ids) -> dict:
    return {_norm_id(v): v for v in valid_ids}


def match_control_id(token: str, lookup: dict) -> str | None:
    return lookup.get(_norm_id(token))


def _rows_to_answers(rows: list[list[str]], lookup: dict) -> list[dict]:
    """Given rows of cells, detect a header and map columns; else positional."""
    if not rows:
        return []
    # Header detection
    col = {"id": 0, "status": 1, "evidence": 2, "department": 3}
    header_used = False
    first = [str(c or "").strip().lower() for c in rows[0]]
    hits = {k: None for k in col}
    for idx, cell in enumerate(first):
        for key, hints in _HEADER_HINTS.items():
            if cell in hints and hits[key] is None:
                hits[key] = idx
    if hits["id"] is not None and hits["status"] is not None:
        col = {k: (v if v is not None else -1) for k, v in hits.items()}
        header_used = True
    body = rows[1:] if header_used else rows

    out = []
    for r in body:
        cells = [str(c).strip() if c is not None else "" for c in r]
        def get(i):
            return cells[i] if 0 <= i < len(cells) else ""
        cid = match_control_id(get(col["id"]), lookup)
        st = normalize_status(get(col["status"]))
        if not cid or not st:
            continue
        out.append({"controlId": cid, "status": st,
                    "evidence": get(col["evidence"]) if col["evidence"] >= 0 else "",
                    "department": get(col["department"]) if col["department"] >= 0 else ""})
    return out


_STRONG_DELIM = re.compile(r"\t|\s*\|\s*|\s*;\s*|\s*,\s*|\s*:\s*|\s{2,}")


def _parse_line(line: str, lookup: dict) -> dict | None:
    line = line.strip()
    if not line:
        return None
    # 1) Strong delimiters (comma / pipe / colon / tab / 2+ spaces)
    parts = [p.strip() for p in _STRONG_DELIM.split(line) if p.strip()]
    if len(parts) >= 2:
        cid = match_control_id(parts[0], lookup)
        st = normalize_status(parts[1])
        if cid and st:
            return {"controlId": cid, "status": st,
                    "evidence": parts[2] if len(parts) > 2 else "",
                    "department": parts[3] if len(parts) > 3 else ""}
    # 2) Whitespace-separated: <control id> <status...> <evidence>
    ws = line.split()
    if not ws:
        return None
    idx = next((i for i, t in enumerate(ws) if match_control_id(t, lookup)), None)
    if idx is None:
        return None
    cid = match_control_id(ws[idx], lookup)
    rest = ws[idx + 1:]
    st, consumed = None, 0
    if len(rest) >= 2 and normalize_status(rest[0] + " " + rest[1]):
        st, consumed = normalize_status(rest[0] + " " + rest[1]), 2
    elif rest and normalize_status(rest[0]):
        st, consumed = normalize_status(rest[0]), 1
    if not st:
        return None
    return {"controlId": cid, "status": st, "evidence": " ".join(rest[consumed:]), "department": ""}


def parse_text(text: str, lookup: dict) -> list[dict]:
    text = (text or "").strip()
    if not text:
        return []
    # JSON?
    if text[0] in "[{":
        try:
            data = json.loads(text)
            items = data.get("assertions", []) if isinstance(data, dict) else data
            out = []
            for a in items or []:
                cid = match_control_id(str(a.get("controlId", "")), lookup)
                st = normalize_status(str(a.get("status", "")))
                if not cid or not st:
                    continue
                dept = (a.get("source", {}) or {}).get("department", "") if isinstance(a.get("source"), dict) else a.get("department", "")
                out.append({"controlId": cid, "status": st,
                            "evidence": str(a.get("evidence", "")), "department": dept})
            return out
        except (json.JSONDecodeError, AttributeError):
            pass
    # Line-by-line, tolerant of delimiters and prose
    out = []
    for line in text.splitlines():
        row = _parse_line(line, lookup)
        if row:
            out.append(row)
    return out


def parse_csv(data: bytes, lookup: dict) -> list[dict]:
    text = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    return _rows_to_answers(rows, lookup)


def parse_xlsx(data: bytes, lookup: dict) -> list[dict]:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows: list[list[str]] = []
    for ws in wb.worksheets:
        for r in ws.iter_rows(values_only=True):
            rows.append([("" if c is None else str(c)) for c in r])
    return _rows_to_answers(rows, lookup)


def parse_docx(data: bytes, lookup: dict) -> list[dict]:
    from docx import Document
    doc = Document(io.BytesIO(data))
    out: list[dict] = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        out.extend(_rows_to_answers(rows, lookup))
    # also scan paragraph text lines
    paras = "\n".join(p.text for p in doc.paragraphs)
    out.extend(parse_text(paras, lookup))
    return _dedupe(out)


def parse_pdf(data: bytes, lookup: dict) -> list[dict]:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return parse_text(text, lookup)


def _dedupe(answers: list[dict]) -> list[dict]:
    seen, out = set(), []
    for a in answers:
        if a["controlId"] in seen:
            continue
        seen.add(a["controlId"])
        out.append(a)
    return out


def parse_upload(filename: str, data: bytes, lookup: dict) -> list[dict]:
    name = (filename or "").lower()
    try:
        if name.endswith(".csv") or name.endswith(".tsv") or name.endswith(".txt"):
            text = data.decode("utf-8", "replace")
            # merge structured CSV rows with free-form line parsing; CSV wins on conflict
            return _dedupe(parse_csv(data, lookup) + parse_text(text, lookup))
        if name.endswith(".xlsx") or name.endswith(".xlsm"):
            return _dedupe(parse_xlsx(data, lookup))
        if name.endswith(".docx"):
            return parse_docx(data, lookup)
        if name.endswith(".pdf"):
            return _dedupe(parse_pdf(data, lookup))
        if name.endswith(".json"):
            return _dedupe(parse_text(data.decode("utf-8", "replace"), lookup))
    except Exception:
        return []
    # unknown extension: try as text
    return _dedupe(parse_text(data.decode("utf-8", "replace"), lookup))
