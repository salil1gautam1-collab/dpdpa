"""AI-assisted import: upload any document, get proposed checkpoint mappings,
review, and apply. Operator-only. The AI only suggests — a human confirms."""
from __future__ import annotations

import uuid

from fastapi import APIRouter, Depends, HTTPException, Request
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from ..audit import record
from ..db import get_db
from ..models import AiSuggestion, Company, QuestionnaireAnswer, Role
from ..services import ai_mapper
from ..services.import_parser import parse_upload, build_id_lookup  # reuse text extraction
from ..services.rulebook_service import latest_rulebook
from ..templating import render
from .helpers import check_csrf, redirect, require

router = APIRouter()

_LABELS = {"aws": "AWS", "azure": "Azure", "intune": "Intune", "gcp": "GCP",
           "adgpo": "AD/GPO", "firewall": "Firewall"}


def _company(request, db, cid):
    p = require(request, db, roles={Role.admin, Role.analyst, Role.cs})
    c = db.get(Company, cid)
    if not c or c.organization_id != p.user.organization_id:
        raise HTTPException(404, "Company not found")
    return p, c


def _extract_text(filename: str, data: bytes) -> str:
    """Pull raw text from the uploaded file (reuses the import parsers' loaders)."""
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for row in t.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if name.endswith(".pdf"):
            import io
            from pypdf import PdfReader
            return "\n".join((pg.extract_text() or "") for pg in PdfReader(io.BytesIO(data)).pages)
        if name.endswith((".xlsx", ".xlsm")):
            import io
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets:
                for r in ws.iter_rows(values_only=True):
                    rows.append(" | ".join("" if c is None else str(c) for c in r))
            return "\n".join(rows)
        return data.decode("utf-8", errors="replace")
    except Exception:
        return data.decode("utf-8", errors="replace")


@router.post("/companies/{cid}/ai-import")
async def ai_import(cid: str, request: Request, db: Session = Depends(get_db)):
    p, c = _company(request, db, cid)
    form = await request.form()
    check_csrf(p, form.get("csrf", ""))
    upload = form.get("file")
    fname = getattr(upload, "filename", "") or ""
    if not upload or not fname:
        return redirect(f"/companies/{cid}", "Choose a document to analyse.", err=True)
    data = await upload.read()
    if len(data) > 15 * 1024 * 1024:
        return redirect(f"/companies/{cid}", "File too large (max 15 MB).", err=True)
    text = _extract_text(fname, data)

    rb = latest_rulebook(db)
    cats = {x["id"]: x["name"] for x in rb["categories"]}
    suggestions, note = ai_mapper.propose_mappings(text, rb["controls"], cats)
    if not suggestions:
        return redirect(f"/companies/{cid}", note, err=True)

    # Replace any prior pending batch for this company
    db.execute(delete(AiSuggestion).where(AiSuggestion.company_id == cid))
    batch = str(uuid.uuid4())
    for sug in suggestions:
        db.add(AiSuggestion(company_id=cid, batch=batch, control_id=sug["controlId"],
                            status=sug["status"], evidence=sug["evidence"],
                            source_quote=sug["sourceQuote"], confidence=sug["confidence"],
                            source_name=fname))
    db.commit()
    record(db, action="ai.import.propose", actor=p.user, target_type="company", target_id=cid,
           ip=getattr(request.state, "client_ip", ""), source=fname, count=len(suggestions))
    return redirect(f"/companies/{cid}/ai-review", note)


@router.get("/companies/{cid}/ai-review")
def ai_review(cid: str, request: Request, db: Session = Depends(get_db)):
    p, c = _company(request, db, cid)
    sugs = list(db.execute(select(AiSuggestion).where(AiSuggestion.company_id == cid)
                           .order_by(AiSuggestion.control_id)).scalars())
    rb = latest_rulebook(db)
    titles = {x["id"]: x["title"] for x in rb["controls"]}
    existing = {a.control_id: a.status for a in db.execute(select(QuestionnaireAnswer).where(
        QuestionnaireAnswer.company_id == cid)).scalars()}
    ok, note = ai_mapper.provider_available()
    return render(request, "ai_review.html", c=c, sugs=sugs, titles=titles,
                  existing=existing, statuses=["COMPLIANT", "PARTIAL", "GAP", "NA", "TBC"],
                  provider_note=note)


@router.post("/companies/{cid}/ai-review/apply")
async def ai_apply(cid: str, request: Request, db: Session = Depends(get_db)):
    p, c = _company(request, db, cid)
    form = await request.form()
    check_csrf(p, form.get("csrf", ""))
    sugs = {s.id: s for s in db.execute(select(AiSuggestion).where(
        AiSuggestion.company_id == cid)).scalars()}
    existing = {a.control_id: a for a in db.execute(select(QuestionnaireAnswer).where(
        QuestionnaireAnswer.company_id == cid)).scalars()}
    applied = 0
    for sid, sug in sugs.items():
        if form.get(f"accept-{sid}") != "1":
            continue
        st = (form.get(f"status-{sid}", "") or sug.status).upper()
        if st not in {"COMPLIANT", "PARTIAL", "GAP", "NA", "TBC"}:
            continue
        ev = (form.get(f"ev-{sid}", "") or sug.evidence).strip()
        if sug.control_id in existing:
            existing[sug.control_id].status = st
            existing[sug.control_id].evidence = ev
            existing[sug.control_id].department = "AI-assisted (reviewed)"
        else:
            db.add(QuestionnaireAnswer(company_id=cid, control_id=sug.control_id, status=st,
                                       evidence=ev, department="AI-assisted (reviewed)"))
        applied += 1
    db.execute(delete(AiSuggestion).where(AiSuggestion.company_id == cid))
    db.commit()
    record(db, action="ai.import.apply", actor=p.user, target_type="company", target_id=cid,
           ip=getattr(request.state, "client_ip", ""), applied=applied)
    return redirect(f"/companies/{cid}", f"Applied {applied} reviewed answer(s) to the questionnaire.")


@router.post("/companies/{cid}/ai-review/discard")
async def ai_discard(cid: str, request: Request, db: Session = Depends(get_db)):
    p, c = _company(request, db, cid)
    form = await request.form()
    check_csrf(p, form.get("csrf", ""))
    db.execute(delete(AiSuggestion).where(AiSuggestion.company_id == cid))
    db.commit()
    return redirect(f"/companies/{cid}", "Discarded the AI suggestions.")
